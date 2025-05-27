import os
import re
from typing import Union
import numpy as np
import pdfplumber
from pdfminer.high_level import extract_text
import mammoth
from PIL import Image, ImageEnhance, ImageFilter
import pytesseract
from docx import Document  


class ResumeTextExtractor:
    """A class for extracting text from resumes in various formats."""

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.supported_formats = ['.pdf', '.docx', '.txt', '.jpg', '.jpeg', '.png']

    def extract(self) -> Union[str, None]:
        """Main method to extract text based on file type."""
        if not os.path.exists(self.file_path):
            return f"Error: File not found at {self.file_path}"

        _, ext = os.path.splitext(self.file_path)
        ext = ext.lower()

        try:
            if ext == '.pdf':
                return self._extract_from_pdf()

            elif ext == '.docx':
                return self._extract_from_docx()

            elif ext == '.txt':
                return self._extract_from_txt()

            elif ext in ('.jpg', '.jpeg', '.png'):
                return self._extract_from_image()

            else:
                return f"Error: Unsupported file type {ext}. Supported formats: {', '.join(self.supported_formats)}"

        except Exception as e:
            return f"Error extracting text from {self.file_path}: {str(e)}"

    def _extract_from_pdf(self) -> str:
        """Extract text from PDF, intelligently handling tables."""
        has_tables = False
        with pdfplumber.open(self.file_path) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                if tables and any(any(row) for row in tables):
                    has_tables = True
                    break

        if has_tables:
            text = ""
            with pdfplumber.open(self.file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text.strip() + "\n"
            return text.strip()
        else:
            return extract_text(self.file_path)

    def _extract_from_docx(self) -> str:
        """Extract text from .docx file, including headers if present."""
        full_text = ""

        # Extract body content with Mammoth
        with open(self.file_path, "rb") as docx_file:
            result = mammoth.extract_raw_text(docx_file)
            full_text = result.value

        # Try to include headers using python-docx
        try:
            doc = Document(self.file_path)
            headers_text = []

            for section in doc.sections:
                for header in [section.header, section.first_page_header, section.even_page_header]:
                    if header:
                        for para in header.paragraphs:
                            text = para.text.strip()
                            if text:
                                headers_text.append(text)

            if headers_text:
                full_text = "\n".join(headers_text) + "\n" + full_text

        except Exception as e:
            print(f"⚠️ Warning: Failed to extract headers from {self.file_path} - {str(e)}")

        return full_text
    
    def _extract_from_docx_new(self) -> str:
        """Extract text from .docx file, including headers, footers, and tables if present."""
        # Default: extract main content using Mammoth

    
        with open(self.file_path, "rb") as docx_file:
            result = mammoth.extract_raw_text(docx_file)
            mammoth_text = result.value.strip()

        try:
            doc = Document(self.file_path)
            headers_text, footers_text, tables_text = [], [], []

            # Collect headers
            for section in doc.sections:
                for header in [section.header, section.first_page_header, section.even_page_header]:
                    if header:
                        for para in header.paragraphs:
                            text = para.text.strip()
                            if text:
                                headers_text.append(text)

            # Collect footers
            for section in doc.sections:
                for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                    if footer:
                        for para in footer.paragraphs:
                            text = para.text.strip()
                            if text:
                                footers_text.append(text)

            # Collect table contents
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        tables_text.append(" | ".join(row_text))

            # If no extra content, just return Mammoth output
            if not (headers_text or footers_text or tables_text):
                return mammoth_text

            # Combine everything
            combined_parts = headers_text + [mammoth_text] + tables_text + footers_text
            return "\n\n".join(combined_parts).strip()

        except Exception as e:
            print(f"⚠️ Warning: Failed to extract headers/footers/tables from {self.file_path} - {str(e)}")
            return mammoth_text

    def _extract_from_txt(self) -> str:
        """Extract and clean text from .txt file."""
        with open(self.file_path, 'r', encoding='utf-8') as file:
            lines = [line.strip() for line in file.readlines() if line.strip()]
        text = "\n".join(lines)
        return re.sub(r'\n{3,}', '\n\n', text)

    def _extract_from_image(self) -> str:
        """
        Extract text from image file with optimized preprocessing for resume documents.
        Applies minimal yet necessary preprocessing techniques to improve OCR accuracy.
        """
        try:
            # Load image
            img = Image.open(self.file_path)
            
            # Convert to RGB if necessary (handles various color modes)
            if img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Apply preprocessing pipeline for better OCR results
            processed_img = self._preprocess_image_for_ocr(img)
            
            # Configure Tesseract for document text (resumes)
            custom_config = r'--oem 3 --psm 6 -c tessedit_char_whitelist=ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789.,!?;:()\'"@+-/%&*[]{}|\\~`#$^_= \n\t'
            
            # Extract text with optimized configuration
            extracted_text = pytesseract.image_to_string(
                processed_img, 
                config=custom_config,
                lang='eng'
            )
            
            return self._post_process_ocr_text(extracted_text)
            
        except Exception as e:
            return f"Error extracting text from image {self.file_path}: {e}"

    def _preprocess_image_for_ocr(self, img: Image.Image) -> Image.Image:
        """
        Apply minimal preprocessing techniques to improve OCR accuracy.
        Based on research showing 25-40% improvement with proper preprocessing.
        """
        # Convert PIL Image to numpy array for processing
        img_array = np.array(img)
        
        # 1. Convert to grayscale (reduces noise, improves processing speed)
        if len(img_array.shape) == 3:
            img = img.convert('L')
        
        # 2. Enhance contrast (improves text-background separation)
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(1.5)  # Moderate contrast boost
        
        # 3. Apply slight sharpening (improves character edge definition)
        img = img.filter(ImageFilter.SHARPEN)
        
        # 4. Resize if image is too small (Tesseract works better with larger images)
        width, height = img.size
        if width < 1000 or height < 1000:
            # Scale up while maintaining aspect ratio
            scale_factor = max(1000 / width, 1000 / height)
            new_width = int(width * scale_factor)
            new_height = int(height * scale_factor)
            img = img.resize((new_width, new_height), Image.LANCZOS)
        
        # 5. Noise reduction (remove small artifacts that confuse OCR)
        img = img.filter(ImageFilter.MedianFilter(size=3))
        
        return img

    def _post_process_ocr_text(self, text: str) -> str:
        """
        Clean up OCR output to improve quality for resume parsing.
        """
        if not text or not text.strip():
            return ""
        
        # Remove excessive whitespace and empty lines
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        # Fix common OCR errors in resumes
        cleaned_text = '\n'.join(lines)
        
        # Common OCR character corrections for resumes
        ocr_corrections = {
            '|': 'I',  # Common OCR error
            '0': 'O',  # In names/words context
            '5': 'S',  # When appropriate
            '1': 'l',  # In word context
            '@': 'a',  # When not in email context
        }
        
        # Apply corrections carefully (only when not in email/phone context)
        
        # Don't correct @ in email addresses
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, cleaned_text)
        
        # Don't correct numbers in phone numbers or dates
        phone_pattern = r'[\+]?[1-9]?[\d\s\-\(\)]{7,15}'
        phones = re.findall(phone_pattern, cleaned_text)
        
        # Apply smart corrections (avoiding emails and phones)
        for old_char, new_char in ocr_corrections.items():
            if old_char == '@' and any(email in cleaned_text for email in emails):
                continue  # Skip @ correction if emails present
            
            # Apply correction in word contexts only
            if old_char in ['0', '5', '1'] and any(phone in cleaned_text for phone in phones):
                continue  # Skip number corrections in phone context
                
            # Apply targeted corrections
            cleaned_text = self._smart_char_replacement(cleaned_text, old_char, new_char)
        
        return cleaned_text

    def _smart_char_replacement(self, text: str, old_char: str, new_char: str) -> str:
        """
        Replace characters only in appropriate contexts (word boundaries).
        """
        # Only replace if the character is surrounded by letters (word context)
        if old_char in ['0', '5', '1', '|']:
            # Replace only when surrounded by letters
            pattern = r'(?<=[A-Za-z])' + re.escape(old_char) + r'(?=[A-Za-z])'
            text = re.sub(pattern, new_char, text)
        
        return text

    def clean_cv_text(self, text: str) -> str:
        """
        Given raw extracted CV text, collapse multiple blank lines into one
        and remove any trailing blank lines at the end.
        """
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # 1) Collapse runs of 2 or more blank lines into exactly one: "\n\n"
        text = re.sub(r'\n\s*\n+', '\n\n', text)
        
        # 2) Strip any leading/trailing whitespace and blank lines
        text = text.strip('\n')
        
        return text